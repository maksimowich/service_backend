from fastapi import FastAPI, UploadFile, File
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores.pgvector import PGVector
from langchain.docstore.document import Document
import os
import sys
import docx


app = FastAPI()

CONNECTION_STRING = "postgresql+psycopg2://syash:2004@localhost:5432/resumes"
COLLECTION_NAME = sys.argv[2]
os.environ["OPENAI_API_KEY"] = sys.argv[1]

embeddings = OpenAIEmbeddings()
STORE = PGVector(
    collection_name=COLLECTION_NAME,
    connection_string=CONNECTION_STRING,
    embedding_function=embeddings,
)

@app.post("/upload_resume/")
async def upload_file(file: UploadFile):
    file_content = file.file.read().decode("utf-8")
    if file:
        STORE.add_documents([Document(page_content=file_content)])
        return {"message": "File uploaded and saved successfully"}
    else:
        return {"message": "No file uploaded"}

@app.get("/score_resume/")
async def resumes_with_score(file: UploadFile):
    if file.filename.endswith('.docx'):
        doc = docx.Document(file.file._file)
        file_content = ""
        for paragraph in doc.paragraphs:
            file_content += paragraph.text + "\n"
    elif file.filename.endswith('.pdf'):
        print(1)
    else:
        file_content = file.file.read().decode("utf-8")

    print(file_content)

    docs_with_score = STORE.similarity_search_with_score(file_content, 5)
    for doc, score in docs_with_score:
        print("-" * 80)
        print("Score: ", score)
        print(doc.page_content)
        print("-" * 80)
    return [(score, doc.page_content) for doc, score in docs_with_score]


def main():
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=30_000)
